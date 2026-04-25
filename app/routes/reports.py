from flask import Blueprint, render_template, send_file, request
from flask_login import login_required, current_user
from app.models import Farm, Sale, SaleEntry, Customer, ChickenBatch, ChickenDaily, FeedUsage, FeedStock
from app import db
import io
from datetime import datetime

reports_bp = Blueprint('reports', __name__, url_prefix='/reports')


@reports_bp.route('/sale/<int:sale_id>')
@login_required
def sale_report(sale_id):
    sale = Sale.query.get_or_404(sale_id)
    farm = Farm.query.filter_by(id=sale.farm_id).first_or_404()
    return render_template('reports/sale.html', sale=sale, farm=farm)


@reports_bp.route('/sale/<int:sale_id>/excel')
@login_required
def export_excel(sale_id):
    sale = Sale.query.get_or_404(sale_id)
    farm = Farm.query.filter_by(id=sale.farm_id).first_or_404()

    lang = request.args.get('lang', 'ta')
    
    t = {
        'title': '🐔 POULTRY FARM SALE REPORT' if lang == 'en' else '🐔 பண்ணை விற்பனை அறிக்கை',
        'farm': 'Farm:' if lang == 'en' else 'பண்ணை:',
        'sale_id': 'Sale ID:' if lang == 'en' else 'விற்பனை குறியீடு:',
        'village': 'Village:' if lang == 'en' else 'கிராமம்:',
        'date': 'Date:' if lang == 'en' else 'தேதி:',
        'phone': 'Phone:' if lang == 'en' else 'தொலைபேசி:',
        'status': 'Status:' if lang == 'en' else 'நிலை:',
        'customer': 'Customer:' if lang == 'en' else 'வாடிக்கையாளர்:',
        'vehicle': 'Vehicle:' if lang == 'en' else 'வாகனம்:',
        'headers': ['S.No', 'Empty Boxes', 'Empty Wt (kg)', 'Birds/Box', 'Load Wt (kg)', 'Total Birds', 'Net Wt (kg)'] if lang == 'en' else 
                   ['வ.எண்', 'வெற்று பெட்டிகள்', 'வெற்று எடை (கிலோ)', 'கோழிகள்/பெட்டி', 'ஏற்றும் எடை (கிலோ)', 'மொத்த கோழிகள்', 'நிகர எடை'],
        'total': 'TOTAL' if lang == 'en' else 'மொத்தம்',
        'net_wt': 'Net Weight' if lang == 'en' else 'நிகர எடை',
        'avg_wt': 'Average Weight/Bird' if lang == 'en' else 'சராசரி எடை',
        'tonnage': 'Tonnage' if lang == 'en' else 'டன்னேஜ்',
    }

    from openpyxl import Workbook
    from openpyxl.styles import Font, PatternFill, Alignment, Border, Side
    from openpyxl.utils import get_column_letter

    wb = Workbook()
    ws = wb.active
    ws.title = 'Sale Report' if lang == 'en' else 'விற்பனை அறிக்கை'

    # Colors
    header_fill = PatternFill('solid', start_color='1a472a')
    sub_fill = PatternFill('solid', start_color='2d6a4f')
    total_fill = PatternFill('solid', start_color='d4edda')
    alt_fill = PatternFill('solid', start_color='f0f7f0')

    header_font = Font(name='Arial', bold=True, color='FFFFFF', size=11)
    sub_font = Font(name='Arial', bold=True, color='FFFFFF', size=10)
    bold_font = Font(name='Arial', bold=True, size=10)
    normal_font = Font(name='Arial', size=10)
    total_font = Font(name='Arial', bold=True, size=10, color='1a472a')

    thin = Side(style='thin', color='AAAAAA')
    border = Border(left=thin, right=thin, top=thin, bottom=thin)

    center = Alignment(horizontal='center', vertical='center')
    left = Alignment(horizontal='left', vertical='center')

    def set_cell(cell, value, font=None, fill=None, align=None, bord=None):
        cell.value = value
        if font:
            cell.font = font
        if fill:
            cell.fill = fill
        if align:
            cell.alignment = align
        if bord:
            cell.border = bord

    # Title
    ws.merge_cells('A1:G1')
    set_cell(ws['A1'], t['title'], Font(name='Arial', bold=True, size=14, color='1a472a'), align=center)
    ws.row_dimensions[1].height = 30

    # Farm & Sale Info
    ws.merge_cells('A3:D3')
    set_cell(ws['A3'], f'{t["farm"]} {farm.name}', bold_font, align=left)
    ws.merge_cells('E3:G3')
    set_cell(ws['E3'], f'{t["sale_id"]} {sale.sale_code}', bold_font, align=left)

    ws.merge_cells('A4:D4')
    set_cell(ws['A4'], f'{t["village"]} {farm.village}, {farm.district}', normal_font, align=left)
    ws.merge_cells('E4:G4')
    set_cell(ws['E4'], f'{t["date"]} {sale.sale_date or "N/A"}', normal_font, align=left)

    ws.merge_cells('A5:D5')
    set_cell(ws['A5'], f'{t["phone"]} {farm.phone}', normal_font, align=left)
    ws.merge_cells('E5:G5')
    
    sale_status_str = sale.status.upper()
    if lang == 'ta':
        sale_status_str = 'முடிந்தது' if sale.status == 'completed' else 'வரைவு'
    set_cell(ws['E5'], f'{t["status"]} {sale_status_str}', Font(name='Arial', bold=True, size=10, color='2d6a4f'), align=left)

    if sale.customer:
        ws.merge_cells('A6:D6')
        set_cell(ws['A6'], f'{t["customer"]} {sale.customer.name}', normal_font, align=left)
        ws.merge_cells('E6:G6')
        set_cell(ws['E6'], f'{t["phone"]} {sale.customer.phone}', normal_font, align=left)
        ws.merge_cells('A7:D7')
        set_cell(ws['A7'], f'{t["vehicle"]} {sale.customer.vehicle_number}', normal_font, align=left)

    # Table headers
    headers = t['headers']
    col_widths = [8, 14, 18, 15, 18, 16, 16]

    row = 9
    for col, (h, w) in enumerate(zip(headers, col_widths), 1):
        cell = ws.cell(row=row, column=col)
        set_cell(cell, h, header_font, header_fill, center, border)
        ws.column_dimensions[get_column_letter(col)].width = w
    ws.row_dimensions[row].height = 22

    # Data rows
    total_empty_boxes = 0
    total_empty_weight = 0
    total_chickens_total = 0
    total_load_weight = 0
    total_net_weight = 0

    for i, entry in enumerate(sale.entries, 1):
        row += 1
        total_chickens = entry.empty_boxes * entry.chickens_per_box
        net_weight = entry.load_weight - entry.empty_weight
        fill = alt_fill if i % 2 == 0 else None

        values = [i, entry.empty_boxes, entry.empty_weight, entry.chickens_per_box,
                  entry.load_weight, total_chickens, round(net_weight, 2)]

        total_empty_boxes += entry.empty_boxes
        total_empty_weight += entry.empty_weight
        total_chickens_total += total_chickens
        total_load_weight += entry.load_weight
        total_net_weight += net_weight

        for col, val in enumerate(values, 1):
            cell = ws.cell(row=row, column=col)
            set_cell(cell, val, normal_font, fill, center, border)

    # Totals row
    row += 1
    totals = [t['total'], total_empty_boxes, round(total_empty_weight, 2), '',
              round(total_load_weight, 2), total_chickens_total, round(total_net_weight, 2)]
    for col, val in enumerate(totals, 1):
        cell = ws.cell(row=row, column=col)
        set_cell(cell, val, total_font, total_fill, center, border)
    ws.row_dimensions[row].height = 20

    # Summary
    row += 2
    avg_weight = total_net_weight / total_chickens_total if total_chickens_total > 0 else 0
    tonnage = total_net_weight / 1000

    summaries = [
        (t['net_wt'], f'{round(total_net_weight, 2)} kg'),
        (t['avg_wt'], f'{round(avg_weight, 3)} kg'),
        (t['tonnage'], f'{round(tonnage, 4)} ton'),
    ]

    for label, value in summaries:
        ws.merge_cells(f'A{row}:C{row}')
        set_cell(ws[f'A{row}'], label, bold_font, sub_fill, left)
        ws[f'A{row}'].font = Font(name='Arial', bold=True, size=10, color='FFFFFF')
        ws.merge_cells(f'D{row}:G{row}')
        set_cell(ws[f'D{row}'], value, Font(name='Arial', bold=True, size=11, color='1a472a'), align=center)
        row += 1

    output = io.BytesIO()
    wb.save(output)
    output.seek(0)

    return send_file(
        output,
        download_name=f'Sale_{sale.sale_code}_{farm.name}.xlsx',
        as_attachment=True,
        mimetype='application/vnd.openxmlformats-officedocument.spreadsheetml.sheet'
    )


@reports_bp.route('/sale/<int:sale_id>/pdf')
@login_required
def export_pdf(sale_id):
    sale = Sale.query.get_or_404(sale_id)
    farm = Farm.query.filter_by(id=sale.farm_id).first_or_404()

    return render_template('reports/sale.html', sale=sale, farm=farm, auto_download=True)


@reports_bp.route('/batch/<int:batch_id>/excel')
@login_required
def export_batch_excel(batch_id):
    lang = request.args.get('lang', 'ta')
    t = {
        'title': 'COMPREHENSIVE BATCH REPORT' if lang == 'en' else 'தொகுதி அறிக்கை',
        'sheet': 'Batch Report' if lang == 'en' else 'தொகுதி அறிக்கை',
        'bname': 'Batch Name:' if lang == 'en' else 'தொகுதி பெயர்:',
        'sdate': 'Start Date:' if lang == 'en' else 'தொடக்க தேதி:',
        't1': ['Day', 'Date', 'Deaths', 'Sold', 'Remaining', 'Feed Used (bags)'] if lang == 'en' else ['நாள்', 'தேதி', 'இறப்பு', 'விற்பனை', 'இருப்பு', 'பயன்படுத்திய தீவனம் (பைகள்)'],
        'tot': 'TOTAL' if lang == 'en' else 'மொத்தம்',
        's_sum': 'SALES SUMMARY' if lang == 'en' else 'விற்பனை சுருக்கம்',
        't2': ['Date', 'Code', 'Customer', 'Birds', 'Weight (kg)'] if lang == 'en' else ['தேதி', 'குறியீடு', 'வாடிக்கையாளர்', 'கோழிகள்', 'எடை (கிலோ)'],
        'fcr': 'FCR ANALYSIS' if lang == 'en' else 'FCR பகுப்பாய்வு',
        'feed': 'Total Feed Consumed (kg)' if lang == 'en' else 'மொத்த தீவனம் (கிலோ)',
        'live': 'Total Live Weight (kg)' if lang == 'en' else 'மொத்த நேரடி எடை (கிலோ)',
    }
    
    batch = ChickenBatch.query.get_or_404(batch_id)
    farm = Farm.query.get(batch.farm_id)
    dailies = ChickenDaily.query.filter_by(batch_id=batch.id).order_by(ChickenDaily.day_number).all()
    usages = FeedUsage.query.filter_by(farm_id=farm.id).order_by(FeedUsage.day_number).all()
    sales = Sale.query.filter_by(farm_id=farm.id, status='completed').filter(Sale.sale_date >= batch.start_date).all()

    import io
    from openpyxl import Workbook
    from openpyxl.styles import Font, PatternFill, Alignment
    
    wb = Workbook()
    
    # Section Style
    header_fill = PatternFill('solid', start_color='1a472a')
    header_font = Font(bold=True, color='FFFFFF')
    center = Alignment(horizontal='center')

    ws = wb.active
    ws.title = t['sheet']
    
    # Title
    ws.merge_cells('A1:E1')
    ws['A1'] = f"{t['title']} - {farm.name}"
    ws['A1'].font = Font(bold=True, size=14)
    ws['A1'].alignment = center
    
    # Batch Info
    ws['A3'] = t['bname']
    ws['B3'] = getattr(batch, 'batch_name', f"Batch #{batch.id}" if lang == 'en' else f"தொகுதி #{batch.id}")
    ws['A4'] = t['sdate']
    ws['B4'] = batch.start_date.strftime('%d/%m/%Y')
    
    # 1. Daily Records Table
    row = 6
    for col, h in enumerate(t['t1'], 1):
        cell = ws.cell(row=row, column=col)
        cell.value = h
        cell.fill = header_fill
        cell.font = header_font
        cell.alignment = center
    
    usage_dict = {u.day_number: u.bags_used for u in usages}
    row += 1
    total_deaths = 0
    total_sold = 0
    total_feed = 0
    
    for d in dailies:
        f_used = usage_dict.get(d.day_number, 0)
        ws.cell(row=row, column=1, value=d.day_number)
        ws.cell(row=row, column=2, value=d.entry_date.strftime('%d/%m/%Y') if d.entry_date else '')
        ws.cell(row=row, column=3, value=d.deaths)
        ws.cell(row=row, column=4, value=d.sold_count)
        ws.cell(row=row, column=5, value=d.remaining)
        ws.cell(row=row, column=6, value=f_used)
        total_deaths += d.deaths
        total_sold += d.sold_count
        total_feed += f_used
        row += 1
    
    # Totals Row
    ws.cell(row=row, column=1, value=t['tot'])
    ws.cell(row=row, column=3, value=total_deaths)
    ws.cell(row=row, column=4, value=total_sold)
    ws.cell(row=row, column=6, value=total_feed)
    for i in range(1, 7):
        ws.cell(row=row, column=i).font = Font(bold=True)
    
    # 2. Sales Summary
    row += 2
    ws.merge_cells(f'A{row}:E{row}')
    ws.cell(row=row, column=1, value=t['s_sum']).font = Font(bold=True)
    row += 1
    for col, h in enumerate(t['t2'], 1):
        ws.cell(row=row, column=col, value=h).fill = header_fill
        ws.cell(row=row, column=col).font = header_font
    
    row += 1
    total_sales_kg = 0
    total_sales_birds = 0
    for s in sales:
        s_birds = sum(e.empty_boxes * e.chickens_per_box for e in s.entries)
        s_kg = sum(e.load_weight - e.empty_weight for e in s.entries)
        ws.cell(row=row, column=1, value=s.sale_date.strftime('%d/%m/%Y') if s.sale_date else '')
        ws.cell(row=row, column=2, value=s.sale_code)
        ws.cell(row=row, column=3, value=s.customer.name if s.customer else '')
        ws.cell(row=row, column=4, value=s_birds)
        ws.cell(row=row, column=5, value=round(s_kg, 2))
        total_sales_kg += s_kg
        total_sales_birds += s_birds
        row += 1
    
    # 3. FCR Analysis
    row += 2
    avg_wt = total_sales_kg / total_sales_birds if total_sales_birds > 0 else 0
    rem_birds = dailies[-1].remaining if dailies else batch.initial_count
    total_live_weight = rem_birds * avg_wt
    feed_kg = total_feed * 50
    fcr = feed_kg / total_live_weight if total_live_weight > 0 else 0
    
    ws.cell(row=row, column=1, value=t['fcr']).font = Font(bold=True)
    row += 1
    ws.cell(row=row, column=1, value=t['feed'])
    ws.cell(row=row, column=2, value=round(feed_kg, 2))
    row += 1
    ws.cell(row=row, column=1, value=t['live'])
    ws.cell(row=row, column=2, value=round(total_live_weight, 2))
    row += 1
    ws.cell(row=row, column=1, value="FCR")
    ws.cell(row=row, column=2, value=round(fcr, 2)).font = Font(bold=True, color='e76f51')

    output = io.BytesIO()
    wb.save(output)
    output.seek(0)
    return send_file(output, download_name=f"Comprehensive_Report_{farm.name}.xlsx", as_attachment=True)


@reports_bp.route('/batch/<int:batch_id>/pdf')
@login_required
def export_batch_pdf(batch_id):
    batch = ChickenBatch.query.get_or_404(batch_id)
    farm = Farm.query.get(batch.farm_id)
    dailies = ChickenDaily.query.filter_by(batch_id=batch.id).order_by(ChickenDaily.day_number).all()
    usages = FeedUsage.query.filter_by(farm_id=farm.id).order_by(FeedUsage.day_number).all()
    sales = Sale.query.filter_by(farm_id=farm.id, status='completed').filter(Sale.sale_date >= batch.start_date).all()
    feed_purchases = FeedStock.query.filter_by(farm_id=farm.id).filter(FeedStock.date >= batch.start_date).order_by(FeedStock.date).all()

    u_dict = {u.day_number: u.bags_used for u in usages}
    p_dict = {}
    for p in feed_purchases:
        p_dict[p.date] = p_dict.get(p.date, 0) + p.bags_added
        
    return render_template('reports/batch_pdf.html', batch=batch, farm=farm, dailies=dailies, usages=usages, sales=sales, u_dict=u_dict, p_dict=p_dict)


@reports_bp.route('/batch/<int:batch_id>/docx')
@login_required
def export_batch_docx(batch_id):
    batch = ChickenBatch.query.get_or_404(batch_id)
    farm = Farm.query.get(batch.farm_id)
    dailies = ChickenDaily.query.filter_by(batch_id=batch.id).order_by(ChickenDaily.day_number).all()
    usages = FeedUsage.query.filter_by(farm_id=farm.id).order_by(FeedUsage.day_number).all()
    sales = Sale.query.filter_by(farm_id=farm.id, status='completed').filter(Sale.sale_date >= batch.start_date).all()

    try:
        from docx import Document
    except ImportError:
        return "python-docx not installed.", 500

    doc = Document()
    doc.add_heading(f'Comprehensive Batch Report - {farm.name}', 0)
    
    # Monitoring
    doc.add_heading('1. Chicken Daily Monitoring', level=1)
    table = doc.add_table(rows=1, cols=5)
    table.style = 'Table Grid'
    hdr = table.rows[0].cells
    hdr[0].text = 'Day'; hdr[1].text = 'Date'; hdr[2].text = 'Deaths'; hdr[3].text = 'Sold'; hdr[4].text = 'Remaining'
    for d in dailies:
        r = table.add_row().cells
        r[0].text = str(d.day_number); r[1].text = d.entry_date.strftime('%d/%m/%Y'); r[2].text = str(d.deaths); r[3].text = str(d.sold_count); r[4].text = str(d.remaining)

    # Feed
    doc.add_heading('2. Feed Daily Usage', level=1)
    ftable = doc.add_table(rows=1, cols=3)
    ftable.style = 'Table Grid'
    fhdr = ftable.rows[0].cells
    fhdr[0].text = 'Day'; fhdr[1].text = 'Date'; fhdr[2].text = 'Bags Used'
    u_dict = {u.day_number: u.bags_used for u in usages}
    total_f = 0
    for d in dailies:
        f = u_dict.get(d.day_number, 0)
        total_f += f
        r = ftable.add_row().cells
        r[0].text = str(d.day_number); r[1].text = d.entry_date.strftime('%d/%m/%Y'); r[2].text = f"{f:.1f}"

    # FCR
    total_sales_kg = sum(sum(e.load_weight - e.empty_weight for e in s.entries) for s in sales)
    total_birds = sum(sum(e.empty_boxes * e.chickens_per_box for e in s.entries) for s in sales)
    avg_wt = total_sales_kg / total_birds if total_birds > 0 else 0
    fcr = (total_f * 50) / (dailies[-1].remaining * avg_wt) if dailies and (dailies[-1].remaining * avg_wt) > 0 else 0
    
    doc.add_heading('3. FCR Analysis', level=1)
    doc.add_paragraph(f'Total Feed Consumed: {total_f * 50:.2f} kg')
    doc.add_paragraph(f'FCR Value: {fcr:.2f}')

    output = io.BytesIO()
    doc.save(output)
    output.seek(0)
    return send_file(output, download_name=f"Comprehensive_Report_{farm.name}.docx", as_attachment=True)
